import io
from docx import Document
from openpyxl import Workbook


def format_duration(hours):
    """
    Converts decimal hours to a human-friendly string (e.g. 2.5 -> '2 hr 30 min')
    """
    if not hours:
        return "0 min"
    
    total_minutes = round(hours * 60)
    
    h = total_minutes // 60
    m = total_minutes % 60
    
    parts = []
    if h > 0:
        parts.append(f"{h} hr")
    if m > 0:
        parts.append(f"{m} min")
        
    if not parts:
        return "0 min"
    return " ".join(parts)


def get_hours_minutes(decimal_hours):
    """
    Converts decimal hours back to hours (int) and minutes (int) snapped to nearest 5 minutes
    """
    if not decimal_hours:
        return 0, 0
    total_mins = round(decimal_hours * 60)
    h = total_mins // 60
    m = total_mins % 60
    m = round(m / 5) * 5
    if m >= 60:
        h += 1
        m = 0
    return h, m


def generate_excel_report(report_data, report_version=None):
    wb = Workbook()
    
    # Sheet 1: Completed Tasks
    ws_completed = wb.active
    ws_completed.title = "Completed Tasks"
    
    # Headers
    ws_completed.append([
        "Task Description",
        "Time Spent",
        "Decimal Hours",
        "AI Assessment (%)",
        "Assessment Reason",
        "Verification Status",
        "Approval Status",
        "Employee Notes",
        "Reasoning",
        "Evidence References"
    ])
    
    comp_tasks = report_data.get("completed_tasks", [])
    for task in comp_tasks:
        if isinstance(task, dict):
            est_hours = task.get("estimated_hours", 0)
            expl = task.get("confidence_explanation", [])
            expl_str = "; ".join(expl) if isinstance(expl, list) else str(expl)
            evidence_ids = task.get("evidence_ids", [])
            evidence_refs = ", ".join(evidence_ids) if evidence_ids else ""
            ws_completed.append([
                task.get("task", ""),
                format_duration(est_hours),
                float(est_hours or 0.0),
                task.get("confidence", 0),
                expl_str,
                task.get("verification_status", "Awaiting Supporting Documents"),
                task.get("approval_status", "Pending Review"),
                task.get("employee_notes", ""),
                task.get("reason", ""),
                evidence_refs
            ])
        
    # Add metadata at the bottom
    ws_completed.append([])
    ws_completed.append(["Metadata:"])
    if report_version:
        ws_completed.append(["Report Version", f"v{report_version}"])
    ws_completed.append(["Latest Conversation Date", report_data.get("latest_conversation_date", "Not Available")])
    ws_completed.append([])
    ws_completed.append(["Review Confirmation Status:"])
    ws_completed.append(["Review Confirmed", "Yes" if report_data.get("review_confirmed", False) else "No"])
    ws_completed.append(["Employee Notes", report_data.get("employee_notes", "")])
        
    # Sheet 2: Pending Tasks
    ws_pending = wb.create_sheet(title="Pending Tasks")
    ws_pending.append(["Task Name"])
    pend_tasks = report_data.get("pending_tasks", [])
    for task in pend_tasks:
        ws_pending.append([str(task)])
        
    # Sheet 3: New Requests
    ws_requests = wb.create_sheet(title="New Requests")
    ws_requests.append(["Task Name"])
    new_reqs = report_data.get("new_requests", [])
    for task in new_reqs:
        ws_requests.append([str(task)])
        
    excel_buffer = io.BytesIO()
    wb.save(excel_buffer)
    excel_buffer.seek(0)
    return excel_buffer.getvalue()


def generate_word_report(project_name, client_name, start_date, end_date, report_data, report_version=None):
    doc = Document()
    import docx
    import re
    
    # Title
    title_p = doc.add_paragraph()
    title_run = title_p.add_run("Project Status Report")
    title_run.font.size = docx.shared.Pt(20)
    title_run.bold = True
    
    # Project Metadata Section
    meta_p = doc.add_paragraph()
    meta_p.add_run("Project: ").bold = True
    meta_p.add_run(f"{project_name}\n")
    meta_p.add_run("Client: ").bold = True
    meta_p.add_run(f"{client_name}\n")
    if start_date and end_date:
        meta_p.add_run("Reporting Period: ").bold = True
        meta_p.add_run(f"{start_date.strftime('%d %b %Y')} to {end_date.strftime('%d %b %Y')}\n")
    else:
        meta_p.add_run("Reporting Period: ").bold = True
        meta_p.add_run("Not Specified\n")
    
    import datetime
    meta_p.add_run("Generated: ").bold = True
    meta_p.add_run(f"{datetime.date.today().strftime('%d %b %Y')}\n")
    if report_version:
        meta_p.add_run("Report Version: ").bold = True
        meta_p.add_run(f"v{report_version}\n")
    
    doc.add_paragraph("-" * 50)  # Simple divider
    
    # Helper to add bold-formatted runs
    def add_formatted_runs(p, text):
        parts = text.split("**")
        is_bold = False
        for part in parts:
            if part:
                run = p.add_run(part)
                run.bold = is_bold
            is_bold = not is_bold
            
    # Completed Tasks Verification & Metadata Section (Part 8)
    doc.add_heading("Completed Tasks Details", level=1)
    comp_tasks = report_data.get("completed_tasks", [])
    if not comp_tasks:
        doc.add_paragraph("No completed tasks identified in this period.")
    else:
        for idx, task in enumerate(comp_tasks):
            if isinstance(task, dict):
                p = doc.add_paragraph()
                p.add_run(f"Task {idx + 1}: {task.get('task', '')}\n").bold = True
                est_hours = task.get('estimated_hours', 0)
                p.add_run(f"- Time Spent: {format_duration(est_hours)}\n")
                
                # AI Assessment
                conf = task.get("confidence", 0)
                p.add_run(f"- AI Assessment: {conf}%\n")
                expl = task.get("confidence_explanation", [])
                if expl:
                    p.add_run("  * Assessment Reasons:\n")
                    for e in expl:
                        p.add_run(f"    - {e}\n")
                        
                # Verification Status & Approval Status
                v_status = task.get("verification_status", "Awaiting Supporting Documents")
                app_status = task.get("approval_status", "Pending Review")
                p.add_run(f"- Verification Status: {v_status}\n")
                p.add_run(f"- Approval Status: {app_status}\n")
                
                # Notes
                notes = task.get("employee_notes", "")
                if notes:
                    p.add_run(f"- Employee Notes: {notes}\n")
                    
                # Evidence References
                evidence_ids = task.get("evidence_ids", [])
                if evidence_ids:
                    p.add_run(f"- Evidence References: {', '.join(evidence_ids)}\n")
                    
    doc.add_paragraph("-" * 50)  # Divider
    
    # Parse client_report markdown
    client_report = report_data.get("client_report", "")
    if client_report:
        lines = client_report.split("\n")
        for line in lines:
            stripped = line.strip()
            if not stripped:
                continue
                
            # Parse headings
            if stripped.startswith("### "):
                doc.add_heading(stripped[4:], level=3)
            elif stripped.startswith("## "):
                doc.add_heading(stripped[3:], level=2)
            elif stripped.startswith("# "):
                doc.add_heading(stripped[2:], level=1)
            # Parse list items
            elif stripped.startswith("- ") or stripped.startswith("* "):
                p = doc.add_paragraph(style='List Bullet')
                add_formatted_runs(p, stripped[2:])
            elif re.match(r'^\d+\.\s', stripped):
                # Numbered list
                parts = stripped.split(".", 1)
                p = doc.add_paragraph(style='List Number')
                add_formatted_runs(p, parts[1].strip())
            else:
                p = doc.add_paragraph()
                add_formatted_runs(p, stripped)
    else:
        # Fallback to manual sections if client_report is empty
        doc.add_heading("Executive Summary", level=1)
        doc.add_paragraph(report_data.get("project_summary", ""))
        
        doc.add_heading("Work Completed Summary", level=1)
        for idx, task in enumerate(comp_tasks):
            if isinstance(task, dict):
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(f"{task.get('task', '')}").bold = True
                p.add_run(f" (Time Spent: {format_duration(task.get('estimated_hours', 0.0))})\n")
                p.add_run(f"Reasoning: {task.get('reason', '')}")
                
        doc.add_heading("Pending Work", level=1)
        for task in report_data.get("pending_tasks", []):
            doc.add_paragraph(str(task), style='List Bullet')
            
        doc.add_heading("Client Requests", level=1)
        for task in report_data.get("new_requests", []):
            doc.add_paragraph(str(task), style='List Bullet')

    # Internal Review Section
    doc.add_paragraph("\n" + "-" * 50)
    doc.add_heading("Internal Review & Verification", level=2)
    confirmed_val = report_data.get("review_confirmed", False)
    notes_val = report_data.get("employee_notes", "")
    doc.add_paragraph(f"Review Confirmed by Employee: {'Yes' if confirmed_val else 'No'}")
    doc.add_paragraph(f"Employee Notes: {notes_val if notes_val else 'None'}")
    
    # Signature
    p_footer = doc.add_paragraph()
    p_footer.add_run("\nPrepared by ProjectMind").italic = True
    
    import io
    word_buffer = io.BytesIO()
    doc.save(word_buffer)
    word_buffer.seek(0)
    return word_buffer.getvalue()


def save_report_to_store(project_name, client_name, start_date, end_date, report_data, messages=None):
    """
    Saves a report to the local persisted_reports/ directory.
    Calculates the next version number for the project.
    """
    import os
    import json
    import time
    from datetime import date
    
    os.makedirs("persisted_reports", exist_ok=True)
    
    # Generate project_id from project_name
    project_id = project_name.lower().strip().replace(" ", "-")
    project_id = "".join(c for c in project_id if c.isalnum() or c == "-")
    if not project_id:
        project_id = "unnamed-project"
        
    # Versioning
    version = 1
    for filename in os.listdir("persisted_reports"):
        if filename.startswith(f"{project_id}_v") and filename.endswith(".json"):
            try:
                parts = filename.split("_v")
                if len(parts) > 1:
                    v_str = parts[1].split(".json")[0]
                    v = int(v_str)
                    if v >= version:
                        version = v + 1
            except Exception:
                pass
                
    report_id = f"rep_{int(time.time())}"
    
    period_str = f"{start_date.strftime('%Y-%m-%d')} to {end_date.strftime('%Y-%m-%d')}" if (start_date and end_date) else ""
    
    from intelligence import compute_project_intelligence
    intel = compute_project_intelligence(messages)
    
    persisted_data = {
        "report_id": report_id,
        "project_id": project_id,
        "project_name": project_name,
        "client_name": client_name,
        "version": version,
        "generated_date": date.today().strftime("%Y-%m-%d"),
        "reporting_period": period_str,
        "start_date": start_date.strftime("%Y-%m-%d") if start_date else "",
        "end_date": end_date.strftime("%Y-%m-%d") if end_date else "",
        "report_data": report_data,
        "messages": messages,
        "project_intelligence_engine": intel
    }
    
    class DateTimeEncoder(json.JSONEncoder):
        def default(self, obj):
            import datetime
            if isinstance(obj, (datetime.datetime, datetime.date)):
                return obj.isoformat()
            return super().default(obj)

    filepath = os.path.join("persisted_reports", f"{project_id}_v{version}.json")
    with open(filepath, "w", encoding="utf-8") as f:
        json.dump(persisted_data, f, indent=2, cls=DateTimeEncoder)
        
    return persisted_data


def update_report_in_store(project_id, version, report_data):
    """
    Updates an existing report's report_data (e.g. after manual edits).
    """
    import os
    import json
    filepath = os.path.join("persisted_reports", f"{project_id}_v{version}.json")
    if os.path.exists(filepath):
        try:
            with open(filepath, "r", encoding="utf-8") as f:
                data = json.load(f)
            class DateTimeEncoder(json.JSONEncoder):
                def default(self, obj):
                    import datetime
                    if isinstance(obj, (datetime.datetime, datetime.date)):
                        return obj.isoformat()
                    return super().default(obj)
            data["report_data"] = report_data
            with open(filepath, "w", encoding="utf-8") as f:
                json.dump(data, f, indent=2, cls=DateTimeEncoder)
            return True
        except Exception as e:
            print(f"Error updating report in store: {e}")
    return False


def list_reports_from_store():
    """
    Lists all reports from the persisted_reports/ store.
    """
    import os
    import json
    reports = []
    if not os.path.exists("persisted_reports"):
        return reports
        
    for filename in os.listdir("persisted_reports"):
        if filename.endswith(".json"):
            filepath = os.path.join("persisted_reports", filename)
            try:
                with open(filepath, "r", encoding="utf-8") as f:
                    data = json.load(f)
                if "report_id" in data and "project_id" in data:
                    import datetime
                    for msg in data.get("messages", []):
                        if isinstance(msg.get("date"), str):
                            try:
                                msg["date"] = datetime.datetime.fromisoformat(msg["date"])
                            except Exception:
                                pass
                    reports.append(data)
            except Exception:
                pass
    return reports


def load_report_from_store(project_id, version):
    """
    Loads and deserializes a specific report from the store.
    """
    import os
    import json
    import datetime
    filepath = os.path.join("persisted_reports", f"{project_id}_v{version}.json")
    if os.path.exists(filepath):
        try:
            with open(filepath, "r", encoding="utf-8") as f:
                data = json.load(f)
            # Convert date strings back to datetimes in messages
            for msg in data.get("messages", []):
                if isinstance(msg.get("date"), str):
                    try:
                        msg["date"] = datetime.datetime.fromisoformat(msg["date"])
                    except Exception:
                        pass
            return data
        except Exception as e:
            print(f"Error loading report from store: {e}")
    return None


def delete_report_from_store(project_id, version):
    """
    Deletes a report file from the store.
    """
    import os
    filepath = os.path.join("persisted_reports", f"{project_id}_v{version}.json")
    if os.path.exists(filepath):
        try:
            os.remove(filepath)
            return True
        except Exception:
            pass
    return False


def extract_text_from_upload(uploaded_file):
    """
    Extracts text content and parses the file based on its extension.
    If the file is a native ProjectMind JSON report, returns the parsed dict and True.
    Otherwise, returns the extracted text and False.
    """
    import json
    filename = uploaded_file.name.lower()
    
    if filename.endswith(".json"):
        try:
            content = uploaded_file.read().decode("utf-8")
            data = json.loads(content)
            if isinstance(data, dict) and "report_id" in data and "report_data" in data:
                import datetime
                for msg in data.get("messages", []):
                    if isinstance(msg.get("date"), str):
                        try:
                            msg["date"] = datetime.datetime.fromisoformat(msg["date"])
                        except Exception:
                            pass
                return data, True
            return content, False
        except Exception as e:
            return f"Error reading JSON file: {str(e)}", False
            
    elif filename.endswith(".docx"):
        try:
            import docx
            doc = docx.Document(uploaded_file)
            text_parts = []
            for p in doc.paragraphs:
                if p.text:
                    text_parts.append(p.text)
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells if cell.text]
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            return "\n".join(text_parts), False
        except Exception as e:
            return f"Error reading DOCX file: {str(e)}", False
            
    elif filename.endswith(".xlsx"):
        try:
            import openpyxl
            wb = openpyxl.load_workbook(uploaded_file, data_only=True)
            text_parts = []
            for sheet in wb.worksheets:
                text_parts.append(f"--- Sheet: {sheet.title} ---")
                for row in sheet.iter_rows(values_only=True):
                    row_vals = [str(val) for val in row if val is not None]
                    if row_vals:
                        text_parts.append(" | ".join(row_vals))
            return "\n".join(text_parts), False
        except Exception as e:
            return f"Error reading XLSX file: {str(e)}", False
            
    else:
        # Default to txt/md or raw text
        try:
            content = uploaded_file.read().decode("utf-8")
            return content, False
        except Exception as e:
            return f"Error reading file: {str(e)}", False


def reconstruct_report_from_text(report_text):
    """
    Calls OpenRouter LLM to reconstruct the project state JSON from raw report text.
    """
    import json
    from config import client, MODEL_NAME
    
    prompt = f"""
You are an expert project manager.
Your task is to analyze the following report text and reconstruct the project state into a structured JSON format.

Report Text:
{report_text}

--------------------------------------

You must extract and return the following fields in the exact JSON format specified below:
{{
    "project_name": "Name of the project (if found, else leave empty)",
    "client_name": "Name of the client (if found, else leave empty)",
    "project_summary": "High-level summary of the report",
    "completed_tasks": [
        {{
            "task": "Task or deliverable name",
            "estimated_hours": 0.0,
            "confidence": 85,
            "reason": "Description or progress details of the task",
            "evidence": [],
            "evidence_ids": []
        }}
    ],
    "pending_tasks": [
        "Description of any pending or ongoing tasks"
    ],
    "new_requests": [
        "Description of any new requests or requirements"
    ],
    "client_report": "The markdown or text version of the report"
}}

Return ONLY valid JSON. Do not include markdown formatting or backticks.
"""
    try:
        response = client.chat.completions.create(
            model=MODEL_NAME,
            messages=[{"role": "user", "content": prompt}],
            temperature=0,
            max_tokens=2500
        )
        raw_response = response.choices[0].message.content
        # Extract JSON if returned with markdown blocks
        if "```json" in raw_response:
            raw_response = raw_response.split("```json")[1].split("```")[0].strip()
        elif "```" in raw_response:
            raw_response = raw_response.split("```")[1].split("```")[0].strip()
        data = json.loads(raw_response.strip())
        return data
    except Exception as e:
        print(f"Error during AI reconstruction: {e}")
        # Return a fallback empty structure
        return {
            "project_name": "",
            "client_name": "",
            "project_summary": "Reconstruction failed. Please check logs.",
            "completed_tasks": [],
            "pending_tasks": [],
            "new_requests": [],
            "client_report": f"Raw content:\\n{report_text}"
        }


def set_active_report_in_store(project_id, version):
    import os
    import json
    os.makedirs("persisted_reports", exist_ok=True)
    filepath = os.path.join("persisted_reports", "active_project.json")
    try:
        with open(filepath, "w", encoding="utf-8") as f:
            json.dump({"project_id": project_id, "version": version}, f)
    except Exception as e:
        print(f"Error setting active report: {e}")


def clear_active_report_in_store():
    import os
    filepath = os.path.join("persisted_reports", "active_project.json")
    if os.path.exists(filepath):
        try:
            os.remove(filepath)
        except Exception as e:
            print(f"Error clearing active report: {e}")


def get_active_report_from_store():
    import os
    import json
    filepath = os.path.join("persisted_reports", "active_project.json")
    if os.path.exists(filepath):
        try:
            with open(filepath, "r", encoding="utf-8") as f:
                data = json.load(f)
            return data.get("project_id"), data.get("version")
        except Exception:
            pass
    return None, None
